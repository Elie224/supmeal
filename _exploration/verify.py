import sys
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
from docx import Document
import re

doc = Document(r'C:\Users\KOURO\Desktop\supmeal\_exploration\Rapport_Projet_Fromagerie_Phases_1_a_7_v2.docx')

# Detect section breaks (page breaks)
body = doc.element.body
page_breaks = 0
explicit_breaks = 0
for p in body.iter():
    if p.tag.endswith('}br'):
        # Check for pageBreakBefore
        if p.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
            explicit_breaks += 1

# Find hardcoded "PAGE X" markers I added in comments (none used)
# Estimate: count headings (each starts a new section visually but not a new page)
# Tables add ~5-6 lines each

# Compute line count approx
total_lines = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        total_lines += 1
        continue
    style = p.style.name
    if style.startswith('Heading 1'):
        total_lines += 2  # heading + spacing
    elif style.startswith('Heading 2'):
        total_lines += 1.5
    elif style.startswith('List'):
        words = len(txt.split())
        total_lines += max(1, words / 12)
    else:
        words = len(txt.split())
        total_lines += max(1, words / 12)
for t in doc.tables:
    rows = len(t.rows)
    total_lines += rows * 1.4 + 2

# US Letter at 11pt with 1.15 line spacing = ~50 lines per page
print(f'Lignes equivalentes: {total_lines:.0f}')
print(f'Pages estimes (50 lignes/page): {total_lines/50:.1f}')
print(f'Pages estimes (45 lignes/page): {total_lines/45:.1f}')

# Print structure summary
print()
print('=== STRUCTURE ===')
section_count = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 1'):
        section_count += 1
        print(f'  Section {section_count}: {p.text}')
print(f'Total Heading 1: {section_count}')
print(f'Total Heading 2: {sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading 2"))}')
print(f'Total tables: {len(doc.tables)}')
