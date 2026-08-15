# -*- coding: utf-8 -*-
import sys
sys.stdout = open(sys.stdout.fileno(), mode="w", encoding="utf-8", buffering=1)
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document()
for section in doc.sections:
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15
def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)
def style_table(tbl):
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    if ri == 0:
                        run.font.size = Pt(10.5)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string("FFFFFF".lstrip("#"))
                    else:
                        run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
            if ri == 0:
                set_cell_bg(cell, "2E74B5")
def add_heading(text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string("#2E74B5".lstrip("#"))
        run.font.bold = True
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string("#2E74B5".lstrip("#"))
        run.font.bold = True
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string("#1F4D78".lstrip("#"))
        run.font.bold = True
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
    return h
def add_para(text, size=11, bold=False, italic=False, color=None, align=None, space_after=8):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#")) if color else None
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p
def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    return p
# === PAGE 1 - COVER + EXECUTIVE SUMMARY ===
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJET FROMAGERIE")
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string("#5A5A5A".lstrip("#"))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Prédiction du risque de contamination du lait")
run.font.name = "Calibri"
run.font.size = Pt(26)
run.font.bold = False
run.font.color.rgb = RGBColor.from_string("#1F3A5F".lstrip("#"))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run("Approche méthodologique en 7 phases")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor.from_string("#5A5A5A".lstrip("#"))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("Document de cadrage — Version 1.0")
run.font.name = "Calibri"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor.from_string("#888888".lstrip("#"))
add_heading("Synthèse exécutive", level=1)
add_para("La fromagerie collecte chaque jour du lait auprès d'une centaine de producteurs, organisé en tournées. Les fichiers transmis (campagnes 2023-2024, 2024-2025, 2025-2026) recensent environ 255 événements de déclassement, représentant plus de 290 000 litres de lait perdu sur trois ans, dont environ 73 % imputables à Salmonella, 20 % à Listeria, et le reste à STEC, antibiotiques et autres pathogènes. L'enjeu économique et sanitaire est réel.", align="justify")
add_para("Ce document décrit une approche en sept phases progressives : unification des données dans une base analytique (Phase 1), mise en place de l'infrastructure Cloud et de la réplication (Phase 2), dashboard opérationnel (Phase 3), analyse exploratoire (Phase 4), préparation du dataset Machine Learning avec protection contre les fuites de données (Phase 5), modélisation prédictive par pathogène (Phase 6), puis industrialisation avec mode shadow et monitoring (Phase 7).", align="justify")
add_para("Trois principes guident toute l'approche : (1) aucun modèle n'est entraîné tant que la base n'est pas validée par les utilisateurs métier ; (2) chaque pathogène — Salmonella, Listeria, STEC — est modélisé séparément car les mécanismes sous-jacents diffèrent ; (3) la métrique d'évaluation principale est métier (Recall@K = part des incidents capturés en contrôlant les K % de livraisons les plus risquées) et non la seule accuracy.", align="justify")
add_heading("Prérequis critiques (à valider avant démarrage)", level=2)
add_bullet("Accès à l'historique complet des livraisons (conformes + non conformes) sur les 3 dernières campagnes.")
add_bullet("Schéma et accès en lecture à l'ERP / logiciel métier producteur des livraisons.")
add_bullet("Validation métier des colonnes prioritaires listées en Phase 1.2.")
add_bullet("Volonté d'investir dans une base analytique Cloud (PostgreSQL) séparée du système de production.")
# === PAGE 2 - PHASE 1 ===
add_heading("Phase 1 — Base de données principale", level=1)
add_para("La première phase construit le socle de tout le projet : une base analytique relationnelle où chaque livraison de lait peut être reconstituée de bout en bout, depuis le producteur jusqu'à l'éventuel déclassement. Cette base devient la source unique des dashboards, de l'analyse exploratoire et des modèles prédictifs.", align="justify")
add_heading("1.1 Données déjà disponibles", level=2)
add_bullet("Incidents / déclassements : campagne, pathogène, volumes déclassés, commentaires, actions correctives (3 classeurs Excel).")
add_bullet("Producteurs : noms, codes sources, identifiants provisoires, variantes orthographiques.")
add_bullet("Dates et tournées : dates reconstruites, codes tournée partiels, traçabilité vers les fichiers sources.")
add_bullet("Contrôle qualité : 255 événements unifiés sur trois campagnes, données brutes conservées pour audit.")
add_heading("1.2 Colonnes prioritaires manquantes", level=2)
add_para("Cinq colonnes sont identifiées comme fortement prédictives mais absentes des fichiers transmis. Elles doivent être acquises auprès de la fromagerie avant tout travail de modélisation :", align="justify")
add_bullet("Température du lait — mesures horodatées à la ferme, à la collecte, au transport, à l'arrivée usine.")
add_bullet("Qualité du lait — pH, acidité, cellules somatiques, flore totale.")
add_bullet("Référence de livraison — identifiant unique permettant de relier chaque volume à un producteur et une tournée.")
add_bullet("Nombre de brebis / taille du lot — au moment de la livraison, pas en valeur annuelle.")
add_bullet("Historique complet des livraisons conformes — sans ce dataset, aucune prédiction probabiliste fiable n'est possible.")
add_heading("1.3 Schéma relationnel cible (7 modules)", level=2)
add_para("Le modèle relationnel s'organise autour de la clé métier livraison_id, qui rattache sans ambiguïté le producteur, la date, le volume, les mesures physiques, la tournée et les résultats laboratoire :", align="justify")
tbl = doc.add_table(rows=8, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
data1 = [("Module", "Rôle"), ("producteurs", "Référentiel unique des producteurs et localisation"), ("livraisons", "Table centrale — une ligne par livraison"), ("analyses_laboratoire", "Résultats Salmonella, Listeria, STEC, antibiotiques"), ("temperatures_lait", "Mesures horodatées de température"), ("qualite_lait", "pH, acidité, cellules somatiques, flore totale"), ("tournees_transport", "Collecte, camion, horaires, durée"), ("incidents_declassements", "Historique des événements déjà présent dans les Excel")]
for i, (a, b) in enumerate(data1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
style_table(tbl)
add_heading("1.4 Résultat attendu", level=2)
add_para("À l'issue de la Phase 1, chaque livraison est reconstituable : producteur → collecte → volume → température → pH / qualité → transport → analyse laboratoire → éventuel déclassement. La base est versionnée, auditée et documentée.", align="justify")
# === PAGE 3 - PHASE 2 ===
add_heading("Phase 2 — Infrastructure Cloud et réplication", level=1)
add_para("La Phase 2 met en place la couche d'ingestion : un PostgreSQL Cloud dédié à l'analytique, alimenté automatiquement par réplication incrémentale depuis l'ERP et les fichiers sources. Les dashboards et modèles ne travaillent jamais directement sur la base de production.", align="justify")
add_heading("2.1 Cartographie des sources", level=2)
add_bullet("Identifier l'ERP ou le logiciel métier utilisé (producteurs, livraisons, qualité, stocks).")
add_bullet("Lister les fichiers Excel, exports CSV, bases laboratoire, et capteurs IoT éventuels.")
add_bullet("Documenter pour chaque source : format, fréquence, clé de liaison, propriétaire métier.")
add_heading("2.2 Mise en place de la réplication", level=2)
add_para("La stack de réplication recommandée s'appuie sur l'existant : dlt (data load tool) pour l'extraction-chargement, PostgreSQL Cloud comme destination analytique. L'outillage déjà en place dans l'entreprise (scripts replicate_*.py) peut être étendu avec un nouveau pipeline dédié à la fromagerie, sans dupliquer l'infrastructure.", align="justify")
add_para("Cinq éléments doivent être mis en place et supervisés :", align="justify")
tbl = doc.add_table(rows=6, cols=2)
data2 = [("Élément", "Finalité"), ("Connexion aux sources", "Identifier accès, formats, clés de liaison"), ("Réplication incrémentale", "Conserver des données cohérentes avec le système métier"), ("Contrôles", "Détecter volumes, dates, IDs, doublons, écarts"), ("Sécurité", "Comptes, droits, sauvegardes, chiffrement"), ("Supervision", "Logs, erreurs, dernier chargement réussi")]
for i, (a, b) in enumerate(data2):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
style_table(tbl)
add_heading("2.3 Résultat attendu", level=2)
add_para("Une base Cloud PostgreSQL alimentée automatiquement, traçable et contrôlée. Les données utiles de l'ERP et des autres systèmes sont disponibles dans un environnement dédié à l'analyse, sans perturber les outils de production.", align="justify")
# === PAGE 4 - PHASE 3 ===
add_heading("Phase 3 — Dashboard opérationnel et validation métier", level=1)
add_para("Le premier dashboard est descriptif, sans aucune couche Machine Learning. Son rôle principal est double : donner aux équipes une vision claire de l'activité et de la qualité du lait, et servir de contrôle fonctionnel de la base avant tout travail de modélisation.", align="justify")
add_heading("3.1 Contenu du dashboard", level=2)
tbl = doc.add_table(rows=8, cols=2)
data3 = [("Bloc", "Indicateurs principaux"), ("Activité", "Volume total livré, nombre de livraisons, producteurs actifs"), ("Déclassements", "Volume déclassé, taux de déclassement, nombre d'incidents"), ("Microbiologie", "Salmonella, Listeria, STEC, antibiotiques / RESA"), ("Producteurs", "Producteurs les plus touchés, récidives, dernier incident"), ("Collecte", "Tournées concernées, volumes et conditions de transport"), ("Qualité du lait", "Température, pH et autres mesures disponibles"), ("Géographie / météo", "Carte des exploitations, température extérieure, pluie, humidité")]
for i, (a, b) in enumerate(data3):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
style_table(tbl)
add_heading("3.2 Filtres et navigation", level=2)
add_para("Le dashboard doit permettre d'analyser les résultats par campagne, producteur, pathogène, tournée, statut conforme / déclassé, et zone géographique lorsque la localisation sera disponible. Cette navigation conditionne la qualité de la validation métier.", align="justify")
add_heading("3.3 Contrôle de cohérence avant le Machine Learning", level=2)
add_para("Avant tout passage à la modélisation, le dashboard sert à détecter des anomalies structurelles : producteurs dupliqués, volumes incohérents, températures aberrantes, analyses mal rattachées. Cette validation manuelle est indispensable — un modèle entraîné sur une base non validée propagera les erreurs sans qu'on les voie.", align="justify")
add_heading("3.4 Résultat attendu", level=2)
add_para("Un outil de pilotage opérationnel et une base validée par les utilisateurs métier. Cette validation est le point de départ de la Phase 4 (analyse exploratoire) puis de la Phase 5 (préparation des données Machine Learning).", align="justify")
# === PAGE 5 - PHASE 4 ===
add_heading("Phase 4 — Analyse exploratoire des données (EDA)", level=1)
add_para("L'EDA compare les livraisons conformes et non conformes pour identifier les facteurs réellement associés aux déclassements. Aucune variable n'est supposée causale a priori : les corrélations sont mesurées avant d'être retenues comme features.", align="justify")
add_heading("4.1 Axes d'analyse", level=2)
add_bullet("Temps : évolution mensuelle, saisonnalité, comparaison inter-campagnes, récidives.")
add_bullet("Producteurs : taux de déclassement, fréquence d'incidents, historique récent.")
add_bullet("Qualité du lait : distribution de la température, du pH, de l'acidité, des cellules somatiques.")
add_bullet("Collecte : tournée, durée de stockage / transport, conditions de collecte.")
add_bullet("Environnement : température extérieure, pluie, humidité issues d'une source météo publique ou de station locale.")
add_bullet("Pathogènes : analyses séparées Salmonella, Listeria, STEC, antibiotiques — pas un agrégat « contamination » global.")
add_heading("4.2 Vérifications statistiques", level=2)
add_bullet("Mesure du taux de valeurs manquantes par variable (certaines colonnes peuvent être remplies à <50 %).")
add_bullet("Mesure du déséquilibre entre livraisons conformes et incidents (typiquement <1 % — à anticiper pour le ML).")
add_bullet("Détection des valeurs aberrantes (volume = 0, température <0 °C, pH hors plage 6,4–6,9).")
add_bullet("Test explicite de la variable « Producteur à risque » : si elle est déjà utilisée en interne, c'est probablement le prédicteur le plus fort a priori — à confirmer statistiquement.")
add_heading("4.3 Résultat attendu", level=2)
add_para("Une liste argumentée de variables candidates pour la Phase 5, avec pour chacune : disponibilité réelle, taux de remplissage, force d'association univariée avec chaque pathogène. Les variables sans signal sont écartées ; les variables prometteuses sont validées pour le dataset Machine Learning.", align="justify")
# === PAGE 6 - PHASE 5 ===
add_heading("Phase 5 — Préparation des données et feature engineering", level=1)
add_para("La granularité de référence est désormais explicite : une ligne du dataset correspond à une livraison. Chaque feature est calculée uniquement à partir d'informations disponibles au moment de la prédiction — c'est la règle anti-fuite de données, qui sépare un modèle qui fonctionne d'un modèle qui triche.", align="justify")
add_heading("5.1 Familles de variables", level=2)
tbl = doc.add_table(rows=6, cols=2)
data4 = [("Famille", "Exemples"), ("Historique producteur", "Incidents 30/90/365 j, jours depuis dernier incident, taux historique"), ("Température", "Température actuelle, maximum, durée hors plage, variation tank → usine"), ("pH / qualité", "pH actuel, moyenne historique, écart au niveau habituel"), ("Météo", "Pluie 7 j, humidité 7 j, température moyenne et maximale récente"), ("Collecte", "Volume, tournée, temps de stockage, durée de transport")]
for i, (a, b) in enumerate(data4):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
style_table(tbl)
add_heading("5.2 Protection contre la fuite de données", level=2)
add_para("Chaque feature est associée à un tag available_at : le moment où la valeur devient connue dans le système. Au moment de la construction du dataset, trois familles de variables sont exclues par construction :", align="justify")
add_bullet("Résultats PCR / VIDAS de l'incident que l'on cherche à prédire.")
add_bullet("Actions correctives déclenchées après l'incident (« reprise pack », « nettoyage griffe », « 2 tuyaux silicone remplacés »).")
add_bullet("Statuts post-livraison (déclassé / conforme) — c'est la cible, jamais une variable.")
add_para("Un test automatique vérifie qu'aucune feature n'a un available_at postérieur à la date de la livraison qu'elle est censée décrire.", align="justify")
add_heading("5.3 Labels et découpage temporel", level=2)
add_bullet("Un label binaire distinct par pathogène : Salmonella (0/1), Listeria (0/1), STEC (0/1), éventuellement antibiotiques.")
add_bullet("Découpage temporel strict : entraînement sur 2023-2024 + 2024-2025, validation et test sur 2025-2026. Pas de mélange aléatoire passé / futur.")
add_bullet("Walk-forward validation : recalibrer sur des fenêtres glissantes pour estimer la stabilité du modèle dans le temps.")
add_heading("5.4 Résultat attendu", level=2)
add_para("Un pipeline reproductible produit automatiquement les tables d'entraînement et de test, avec traçabilité des transformations et protection vérifiée contre le data leakage. Les features sont documentées (dictionnaire de données).", align="justify")
# === PAGE 7 - PHASE 6 ===
add_heading("Phase 6 — Modèles prédictifs", level=1)
add_para("La Phase 6 construit et compare des modèles de classification binaire, un par pathogène, avec deux familles complémentaires : une baseline interprétable et un modèle tabulaire plus performant. Les modèles sont entraînés sur les données historiques et évalués sur la campagne 2025-2026 mise de côté.", align="justify")
add_heading("6.1 Approche par pathogène", level=2)
add_para("Les mécanismes à l'œuvre ne sont pas les mêmes : Salmonella est souvent liée à l'environnement de traite et à l'hygiène, Listeria au biofilm et aux équipements (résistance au froid), STEC à la flore digestive et à l'alimentation. Un modèle global « contaminé » mélangerait des signaux antagonistes et perdrait en performance.", align="justify")
add_heading("6.2 Modèles testés", level=2)
tbl = doc.add_table(rows=4, cols=2)
data5 = [("Niveau", "Approche"), ("Baseline", "Régression logistique — référence simple, coefficients directement interprétables"), ("Modèle avancé", "CatBoost (catégorielles natives, robuste aux valeurs manquantes)"), ("Cibles", "P(Salmonella), P(Listeria), P(STEC), éventuellement antibiotiques")]
for i, (a, b) in enumerate(data5):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
style_table(tbl)
add_heading("6.3 Métriques", level=2)
add_para("L'accuracy est inadaptée ici : avec ~99 % de livraisons conformes, un modèle qui répondrait toujours « non contaminé » obtiendrait 99 % d'accuracy sans rien détecter. Les métriques retenues :", align="justify")
add_bullet("Recall et Precision : part des incidents détectés et pertinence des alertes.")
add_bullet("PR-AUC : performance globale sur classes déséquilibrées.")
add_bullet("Recall@5/10/20 % : part des incidents qui seraient capturés si le laboratoire ne contrôle renforcé que les K % de livraisons les plus risquées. C'est la métrique métier centrale.")
add_bullet("Calibration : un risque annoncé de 20 % doit correspondre à ~20 % de cas positifs réels. Sans calibration, le score est inutilisable opérationnellement.")
add_heading("6.4 Explicabilité et validation", level=2)
add_para("Chaque score est accompagné de ses facteurs principaux (historique récent, température, pH, météo, tournée). Les résultats sont présentés aux équipes qualité et laboratoire pour validation. Un modèle qui détecte correctement des incidents pour de mauvaises raisons n'est pas déployable.", align="justify")
add_heading("6.5 Résultat attendu", level=2)
add_para("Des modèles validés sur la campagne 2025-2026, des probabilités calibrées et des seuils de décision testables opérationnellement (ex. : « alerter si P(Salmonella) > 12 % et tournée à risque »).", align="justify")
# === PAGE 8 - PHASE 7 ===
add_heading("Phase 7 — Dashboard prédictif et industrialisation", level=1)
add_para("La dernière phase transforme les modèles en système opérationnel : le dashboard descriptif de la Phase 3 est enrichi d'une couche prédictive, et le flux complet est industrialisé pour produire des scores à chaque nouvelle livraison.", align="justify")
add_heading("7.1 Couche prédictive du dashboard", level=2)
tbl = doc.add_table(rows=5, cols=2)
data6 = [("Information", "Affichage"), ("Contexte livraison", "Producteur, date, volume, tournée, température, pH"), ("Probabilités", "P(Salmonella), P(Listeria), P(STEC)"), ("Priorité", "Faible / Surveillance / Contrôle renforcé"), ("Explication", "Facteurs principaux ayant contribué au score")]
for i, (a, b) in enumerate(data6):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
style_table(tbl)
add_heading("7.2 Industrialisation du flux", level=2)
add_para("Le pipeline complet suit le schéma suivant :", align="justify")
add_para("Base Cloud → pipeline de features → modèles ML → service de prédiction → dashboard prédictif.", align="center", italic=True)
add_bullet("Automatiser le calcul des scores à l'arrivée de nouvelles données (job quotidien ou temps réel selon le besoin).")
add_bullet("Journaliser les versions de modèle et toutes les prédictions (audit, debugging, ré-entraînement).")
add_bullet("Sécuriser les accès et superviser les erreurs de pipeline (alertes en cas d'échec).")
add_heading("7.3 Déploiement progressif et monitoring", level=2)
add_para("Le déploiement suit une trajectoire prudente, indispensable pour un sujet à enjeu sanitaire :", align="justify")
add_bullet("Étape 1 — Shadow mode : les scores sont calculés et affichés, sans modifier les décisions métier. Période d'observation (4 à 8 semaines).")
add_bullet("Étape 2 — Utilisation consultative : les équipes labo utilisent le score comme une alerte parmi d'autres, en complément de leur jugement.")
add_bullet("Étape 3 — Utilisation opérationnelle : le score pilote une partie des contrôles (top K %), avec validation sur les métriques Recall@K et calibration.")
add_bullet("Surveillance continue : dérive des données, dérive du modèle, recalibration automatique, ré-entraînement périodique (trimestriel ou sur signal).")
add_heading("7.4 Résultat attendu", level=2)
add_para("Un système opérationnel de bout en bout : données synchronisées, dashboards, scores de risque, suivi des performances, et processus documenté de maintenance du modèle. Le score est utilisé par le laboratoire pour prioriser ses contrôles.", align="justify")
# === PAGE 9 - RISQUES ===
add_heading("Risques, prérequis et dépendances", level=1)
add_heading("Prérequis métier (à sécuriser en amont)", level=2)
add_bullet("Accès à l'historique complet des livraisons, y compris les conformes, sur 3 ans. Sans ce dataset, la Phase 6 ne peut pas produire de probabilités calibrées.")
add_bullet("Identification d'un référent métier côté fromagerie (qualité, laboratoire, collecte) pour les validations des Phases 1, 3 et 6.")
add_bullet("Clarification RGPD sur les données producteurs (personnes physiques vs entités morales GAEC).")
add_bullet("Validation de la stratégie Cloud (hébergeur, localisation des données, sauvegardes).")
add_heading("Risques techniques et opérationnels", level=2)
add_bullet("Données manquantes : certaines colonnes clés (température, pH, météo) peuvent n'être disponibles que sur une partie de l'historique. À mesurer en Phase 1.")
add_bullet("Déséquilibre des classes : <1 % de livraisons contaminées en moyenne. Modèles et métriques doivent être adaptés.")
add_bullet("Dérive temporelle : les schémas de contamination changent (nouveau pathogène, nouvelle pratique d'élevage). Ré-entraînement périodique indispensable.")
add_bullet("Fuite de données accidentelle : une feature temporelle mal construite peut « voir » le futur. Tests automatiques en Phase 5.")
add_bullet("Adoption métier : un score non utilisé n'a pas de valeur. Le shadow mode de la Phase 7 est le sas d'acceptation.")
add_heading("Indicateurs de succès globaux", level=2)
add_bullet("Base de données unifiée, versionnée, auditée (Phase 1).")
add_bullet("Dashboard opérationnel validé par les équipes métier (Phase 3).")
add_bullet("Recall@10 % ≥ 50 % sur Salmonella en validation temporelle (Phase 6).")
add_bullet("Modèle en shadow mode pendant ≥ 4 semaines avant toute utilisation opérationnelle (Phase 7).")
# === PAGE 10 - ANNEXES ===
add_heading("Annexes", level=1)
add_heading("Annexe A — Glossaire des termes métier", level=2)
add_bullet("Déclassement : opération consistant à retirer un lot de lait de la fabrication fromagère en raison d'un risque sanitaire.")
add_bullet("Tournée : circuit de collecte quotidien (T1-1, T2-1, etc.) regroupant un sous-ensemble de producteurs.")
add_bullet("PCR / VIDAS : méthodes d'analyse laboratoire pour détecter l'ADN (PCR) ou les antigènes (VIDAS) d'un pathogène.")
add_bullet("RESA rapide : test rapide de dépistage des résidus d'antibiotiques dans le lait.")
add_bullet("Producteur à risque : producteur ayant présenté au moins un incident récent, suivi de façon renforcée.")
add_heading("Annexe B — Stack technique recommandée", level=2)
add_bullet("Stockage : PostgreSQL Cloud (couche analytique), isolation du système de production.")
add_bullet("Réplication : dlt (déjà en place), scripts replicate_*.py à étendre.")
add_bullet("Transformation : dbt pour la couche analytique (tests, documentation, versionnage).")
add_bullet("Dashboards : Streamlit pour itérer vite, Metabase pour la diffusion.")
add_bullet("Machine Learning : scikit-learn (baseline) → CatBoost (production).")
add_bullet("Orchestration : Prefect ou Dagster (Python, simples à maintenir).")
add_bullet("MLOps : MLflow pour le tracking, Evidently pour la détection de dérive.")
add_heading("Annexe C — Volumes observés dans les fichiers sources", level=2)
tbl = doc.add_table(rows=5, cols=5)
data7 = [("Campagne", "Salmonella (L)", "Listeria (L)", "STEC (L)", "Total (L)"), ("2023-2024", "139 000", "21 018", "1 244", "161 262"), ("2024-2025", "69 135", "24 469", "1 049", "94 653"), ("2025-2026", "52 354", "7 307", "2 892", "64 048"), ("Total 3 ans", "260 489", "52 794", "5 185", "319 963")]
for i, row_data in enumerate(data7):
    for j, val in enumerate(row_data):
        tbl.rows[i].cells[j].text = val
style_table(tbl)
add_para("Source : agrégats des classeurs Excel RECAP LAITS DECLASSES (campagnes 2023-2024, 2024-2025, 2025-2026).", italic=True, size=9, color="#666666")
add_para("Note : ces totaux reflètent les valeurs agrégées en fin de colonne « total lait déclassé tous pathogènes confondus ». La somme des événements individuels donne un ordre de grandeur cohérent (~290 k L), avec Salmonella qui représente environ 73 % du volume renseigné.", align="justify")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("— Fin du document —")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor.from_string("#888888".lstrip("#"))
out_dir = r"C:\Users\KOURO\Desktop\supmeal\_exploration"
os.makedirs(out_dir, exist_ok=True)
out = os.path.join(out_dir, "Rapport_Projet_Fromagerie_Phases_1_a_7_v2.docx")
doc.save(out)
print("Saved:", out)
print("Size:", os.path.getsize(out), "bytes")


