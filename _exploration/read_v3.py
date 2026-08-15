import sys
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
import docx
d = docx.Document(r'C:\Users\KOURO\Downloads\Rapport_Projet_Fromagerie_v3_corrige.docx')
print(f'Paragraphes: {len(d.paragraphs)}')
print(f'Tables: {len(d.tables)}')
print()
for i, p in enumerate(d.paragraphs):
    style = p.style.name if p.style else 'Normal'
    txt = p.text.strip()
    if txt:
        print(f'[{style[:14]:<14}] {txt[:200]}')
print()
print('=== TABLES ===')
for ti, t in enumerate(d.tables):
    print(f'--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---')
    for ri, r in enumerate(t.rows[:15]):
        cells = [c.text.strip()[:90] for c in r.cells]
        print(f'  R{ri}: {cells}')
