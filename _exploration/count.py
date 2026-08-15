import sys
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
from docx import Document
doc = Document(r'C:\Users\KOURO\Desktop\supmeal\_exploration\Rapport_Projet_Fromagerie_Phases_1_a_7_v2.docx')
total_words = 0
total_chars = 0
heading_count = 0
bullet_count = 0
para_count = 0
table_count = len(doc.tables)
print(f'Tables: {table_count}')
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    para_count += 1
    if p.style.name.startswith('Heading') or p.style.name.startswith('Title'):
        heading_count += 1
    elif p.style.name.startswith('List'):
        bullet_count += 1
    words = len(txt.split())
    total_words += words
    total_chars += len(txt)
print(f'Paragraphes non vides: {para_count}')
print(f'  Headings: {heading_count}')
print(f'  Bullets/numbered: {bullet_count}')
print(f'Total mots: {total_words}')
print(f'Total caracteres: {total_chars}')
print(f'Estimation pages (350 mots/page): {total_words/350:.1f}')
print(f'Estimation pages (450 mots/page): {total_words/450:.1f}')
