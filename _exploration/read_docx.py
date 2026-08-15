import docx
import sys
import os
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

d = docx.Document(r'C:\Users\KOURO\Downloads\Rapport_Projet_Fromagerie_Phases_1_a_7.docx')
print(f'Paragraphes: {len(d.paragraphs)}')
print(f'Tables: {len(d.tables)}')
print()
for i, p in enumerate(d.paragraphs):
    style = p.style.name if p.style else 'Normal'
    txt = p.text.strip()
    if txt:
        print(f'[{style}] {txt}')
print()
print('=== TABLES ===')
for ti, t in enumerate(d.tables):
    print(f'--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---')
    for ri, r in enumerate(t.rows[:25]):
        cells = [c.text.strip()[:100] for c in r.cells]
        print(f'  R{ri}: {cells}')
